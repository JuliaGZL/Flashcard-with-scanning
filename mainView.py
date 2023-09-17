# Let FlashcardApp inherit Database class
from database import Database
# Let FlashcardApp inherit Scanner class
from scanner import Scanner

# Import Tkinter and its children, for building GUI
import tkinter as tk
# Import Tkinter messagebox, for showing message boxes to the user
from tkinter import messagebox
# Import Tkinter ttk, for more visually appealing widgets
from tkinter import ttk
# Import customtkinter, for additionally visually appealing widgets
import customtkinter as ctk
# Import Pillow to display image on tk.Canvas
from PIL import Image, ImageTk
# Import my customized widgets
from view_widgets import WordDelWidget, WordEntryWidget, CardSetWidget, TextListWidget
# Import Card and Cardset classes to be contained
from cards import Card, CardSet
# Import gc to allow string cleansing by regular expressions
import gc
# Import List from typing to indicate the return type of functions
from typing import List, Dict
# Import datetime to get current date and compute dates
from datetime import date, timedelta
# Import numpy for transfering arrays to np.arrays
import numpy as np
# Import pandas for reading and writing excel files
import pandas as pd
# Import python-docx for reading and writing word files
import docx
# Import os to get current file directory
import os
# Import filedialog to let user select directories
from customtkinter import filedialog
# Import cv2 for image processing
import cv2

# Default color themes include blue, green, or dark-blue.
ctk.set_default_color_theme("blue")
ctk.set_appearance_mode("light")


class HomeFrame(ctk.CTkFrame):
    def __init__(self, master: tk.Tk, cardsets: List[CardSet], *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Home")
        self.master.geometry("230x380")
        # self.master.resizable(False, False)

        # Data
        # self.cardsets only used before showing data
        self.cardsets = cardsets
        self.sort_default()
        self.select_idx = None
        self.today = date.today().isoformat()

        # Widgets
        # Card sets
        self.cardsets_viewer = ctk.CTkScrollableFrame(self)
        self.cardset_widgets = []
        for i, cardset in enumerate(self.cardsets):
            self.cardset_widgets.append(CardSetWidget(
                self.cardsets_viewer, cardset, i))
            self.cardset_widgets[-1].pack(padx=5, pady=7)

        # Buttons initialization
        self.btns = ctk.CTkFrame(self)
        self.btn_open = ctk.CTkButton(
            self.btns, text="Open", command=self.open_set, width=200, height=20)
        self.btn_add = ctk.CTkButton(self.btns, text="Add set",
                                     command=self.add_set, width=200, height=20)
        self.btn_delete = ctk.CTkButton(
            self.btns, text="Delete", command=self.delete_set, width=200, height=20)
        self.btn_export = ctk.CTkButton(
            self.btns, text="Export", command=self.export, width=200, height=20)
        # Buttons packing
        self.btn_open.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
        self.btn_add.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
        self.btn_delete.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
        self.btn_export.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)

        # Packing
        self.cardsets_viewer.pack(side=ctk.TOP, padx=5, pady=5)
        self.btns.pack(side=ctk.BOTTOM, padx=5, pady=15)

    # Sets self.select_idx and updates color of border
    def click_cardset(self, event, idx: int) -> None:
        if self.select_idx != None:
            # If changes current selection
            if self.select_idx != idx:
                if self.cardset_widgets[self.select_idx].await_date <= self.today:
                    self.cardset_widgets[self.select_idx].change_color("blue")
                else:
                    self.cardset_widgets[self.select_idx].change_color(
                        "default")
        self.cardset_widgets[idx].change_color("green")
        self.select_idx = idx

    # Opens selected set
    def open_set(self) -> None:
        # Check whether the cardsetView is selected
        if self.select_idx != None:
            self.master.open_frame(
                CardFrame, self.cardset_widgets[self.select_idx].cardset_id)
        else:
            messagebox.showinfo("Error", "Please select a card set.")

    # Shows AddChoiceFrame, the window that gives choices to add a card set
    def add_set(self) -> None:
        self.master.open_frame(AddChoiceFrame)

    # Deletes the selected card set
    def delete_set(self) -> None:
        if self.select_idx != None:
            self.master.delete_cardset(
                self.cardset_widgets[self.select_idx].cardset_id)
        else:
            messagebox.showinfo("Error", "Please select a card set.")
        self.cardsets = self.master.get_all_cardsets()
        self.master.refresh_frame(HomeFrame, self.cardsets)

    # Opens ExportFrame
    def export(self) -> None:
        # Get directory
        folder_path = self._get_export_path()
        # If user selects a folder
        if folder_path != "":
            if self.select_idx != None:
                cardset_id = self.cardset_widgets[self.select_idx].cardset_id
                self.master.open_frame(ExportFrame, cardset_id, folder_path)
        else:
            messagebox.showinfo("Error", "Please select a card set.")
    
    # Opens exporting cardset page
    def export_cardset(self) -> None:
        folder_path = self._get_export_path()
        if folder_path != "":
            if self.select_idx != None:
                cardset = self.master.get_set_cards(
                    self.cardset_widgets[self.select_idx].cardset, self.cardset_widgets[self.select_idx].cardset_id)
                folder_path = os.path.join(
                    folder_path, f"{self.cardset_widgets[self.select_idx].title}.xlsx")
                print(f"Folder path: {folder_path}")
                terms = []
                defs = []
                current = cardset.head
                while current.next != None:
                    terms.append(current.term)
                    defs.append(current.definition)
                    current = current.next
                terms.append(current.term)
                defs.append(current.definition)
                df = pd.DataFrame()
                df["terms"] = np.array(terms)
                df["defs"] = np.array(defs)
                df.to_excel(folder_path, sheet_name="Cards")
                messagebox.showinfo("Success", "Finished exporting.")
            else:
                messagebox.showinfo("Error", "Please select a card set.")
        # Else the process just ends

    # Returns exporting folder path
    def _get_export_path(self) -> str:
        current_path = os.getcwd()
        folder_path = filedialog.askdirectory(
            initialdir=current_path, title="Folder of exported file")
        return folder_path

    # Sorts alphabetically sorted card sets by merge sort
    def _sort_alphabetical(self, cardsets: List[CardSet]) -> List[CardSet]:
        if len(cardsets) > 1:
            mid = len(cardsets) // 2
            # Divide the array into two parts
            left = cardsets[:mid]
            right = cardsets[mid:]
            left = self._sort_alphabetical(left)
            right = self._sort_alphabetical(right)

            # Merge
            l = r = current = 0
            while l < len(left) and r < len(right):
                if left[l].title < right[r].title:
                    cardsets[current] = left[l]
                    l += 1
                else:
                    cardsets[current] = right[r]
                    r += 1
                current += 1
            while l < len(left):
                cardsets[current] = left[l]
                l += 1
                current += 1
            while r < len(right):
                cardsets[current] = right[r]
                r += 1
                current += 1
        # Does not change the cardsets list if base case len(cardsets) == 1
        return cardsets

    # First sorts blue card sets alphabetically, then sorts rest alphabetiall sorted card sets by merge sort
    def sort_default(self) -> None:
        blue_sets = []
        black_sets = []
        for set in self.cardsets:
            if set.awaiting:
                blue_sets.append(set)
            else:
                black_sets.append(set)
        blue_sets = self._sort_alphabetical(blue_sets)
        black_sets = self._sort_alphabetical(black_sets)
        for i in range(len(blue_sets)):
            self.cardsets[i] = blue_sets[i]
        for i in range(len(black_sets)):
            self.cardsets[len(blue_sets)+i] = black_sets[i]

class ExportFrame(ctk.CTkFrame):
    def __init__(self, master, cardset_id: int, folder_path: str, *args):
        # Initialize window
        super().__init__(master, *args)
        self.master.geometry("300x200")
        self.master.title("Export by...")
        
        # Data
        self.cardset_id = cardset_id
        print("Cardset id", self.cardset_id)
        self.folder_path = folder_path
        self.cardset = self.master.get_cardset(self.cardset_id)
        self.cardset = self.master.get_set_cards(self.cardset, self.cardset_id)
        print("Cardset")
        print(self.cardset)
        
        # Widgets - initialization
        self.excel_btn = ctk.CTkButton(self, text="Excel", command=self.export_excel)
        self.txt_btn = ctk.CTkButton(self, text="txt", command=self.export_txt)
        # Widgets - packing
        self.excel_btn.pack(side=ctk.TOP, padx=5, pady=10)
        self.txt_btn.pack(side=ctk.TOP, padx=5, pady=10)
 
    # Exports an excel file    
    def export_excel(self) -> None:
        folder_path = os.path.join(
                    self.folder_path, f"{self.cardset.title}.xlsx")
        terms = []
        defs = []
        current = self.cardset.head
        while current.next != None:
            terms.append(current.term)
            defs.append(current.definition)
            current = current.next
        terms.append(current.term)
        defs.append(current.definition)
        df = pd.DataFrame()
        df["terms"] = np.array(terms)
        df["defs"] = np.array(defs)
        df.to_excel(folder_path, sheet_name="Cards")
        messagebox.showinfo("Success", "Finished exporting.")
        self.master.to_home()

    # Exports a txt file
    def export_txt(self) -> None:
        folder_path = os.path.join(
                    self.folder_path, f"{self.cardset.title}.txt")
        with open(folder_path, 'w') as f:
            current = self.cardset.head
            while current.next != None:
                text = f"{current.term};; {current.definition}\n"
                f.write(text)
                current = current.next
            text = f"{current.term};; {current.definition}"
            f.write(text)
        messagebox.showinfo("Success", "Finished exporting.")
        self.master.to_home()

class AddChoiceFrame(ctk.CTkFrame):
    def __init__(self, master: tk.Tk, *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Add by...")
        self.master.geometry("200x250")

        # Initialize widgets
        # Entry initialization
        self.title_frame = ctk.CTkFrame(self)
        self.lbl_title = ctk.CTkLabel(
            self.title_frame, text="Input title of cardset:")
        self.entry_title = ctk.CTkEntry(self.title_frame)
        # Entry packing
        self.lbl_title.pack(padx=5, pady=5)
        self.entry_title.pack(padx=5, pady=5)
        # Buttons initialization
        self.btns = ctk.CTkFrame(self)
        self.btn_manual = ctk.CTkButton(
            self.btns, text="Manual", command=self.manual_add)
        self.btn_auto = ctk.CTkButton(
            self.btns, text="Image", command=self.image_add)
        self.btn_download = ctk.CTkButton(
            self.btns, text="Document", command=self.document_add)
        # Buttons packing
        self.btn_manual.pack(padx=5, pady=5)
        self.btn_auto.pack(padx=5, pady=5)
        self.btn_download.pack(padx=5, pady=5)

        # Packing widgets
        self.title_frame.pack(side=ctk.TOP, padx=5, pady=7)
        self.btns.pack(side=ctk.TOP, padx=5, pady=7)

    # Opens ManualAddSetFrame
    def manual_add(self) -> None:
        title = self.entry_title.get()
        if self.master.is_empty_string(title):
            messagebox.showinfo("Error", "Empty title!")
        elif len(title) >= 50:
            messagebox.showinfo("Error", "Title is too long!")
        else:
            self.master.open_frame(ManualAddSetFrame, title)

    # Opens ImageAddSetFrame
    def image_add(self) -> None:
        title = self.entry_title.get()
        if self.master.is_empty_string(title):
            messagebox.showinfo("Error", "Empty title!")
        elif len(title) >= 50:
            messagebox.showinfo("Error", "Title is too long!")
        else:
            filetypes = [
                ("Images", ".jpg .png .jpeg .bmp"),
            ]
            paths = filedialog.askopenfilenames(
                title="Select image(s)", filetypes=filetypes)
            if paths != "":
                self.master.open_frame(ImageAddSetFrame, title, paths)

    # Opens DocumentAddSetFrame
    def document_add(self) -> None:
        title = self.entry_title.get()
        if self.master.is_empty_string(title):
            messagebox.showinfo("Error", "Empty title!")
        elif len(title) >= 50:
            messagebox.showinfo("Error", "Title is too long!")
        else:
            self.master.open_frame(DocumentChoiceFrame, title)


class ManualAddSetFrame(ctk.CTkFrame):
    def __init__(self, master: tk.Frame, title: str, *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Adding by Manual")
        self.master.geometry("700x1000")

        # Initialize data
        self.today = date.today()
        tmr = self.today + timedelta(days=1)
        self.cardset = CardSet(id=-1, title=title)
        self._current_idx = 0
        self.cardset.insert(
            Card(id=-1, term="", definition=""))
        self.current = self.cardset.head

        # Initialize widgets
        font = ctk.CTkFont(family="Helvetica", size=20)
        self.term_frame = ctk.CTkFrame(self)
        # Term frame initialization
        self.term_lbl = ctk.CTkLabel(
            self.term_frame, text="Term {}".format(self._current_idx + 1))
        self.term_txt = ctk.CTkTextbox(
            self.term_frame, height=150, width=680, font=font)
        # Term frame packing
        self.term_lbl.pack(side=ctk.TOP, padx=5, pady=5)
        self.term_txt.pack(side=ctk.BOTTOM, padx=5, pady=5)
        # Def frame initialization
        self.def_frame = ctk.CTkFrame(self)
        self.def_lbl = ctk.CTkLabel(
            self.def_frame, text="Definition {}".format(self._current_idx + 1))
        self.def_txt = ctk.CTkTextbox(
            self.def_frame, height=200, width=680, font=font)
        self.def_lbl.pack(side=ctk.TOP, padx=5, pady=5)
        self.def_txt.pack(side=ctk.BOTTOM, padx=5, pady=5)

        # Packing the frames above
        self.term_frame.pack(side=ctk.TOP, padx=5, pady=7)
        self.def_frame.pack(side=ctk.TOP, padx=5, pady=10)

        # Buttons initialization
        self.btn_y1 = ctk.CTkFrame(self)
        self.btn_y2 = ctk.CTkFrame(self)
        self.btns_flip = ctk.CTkFrame(self.btn_y1)
        # btn_back and btn_next
        self.btn_next = ctk.CTkButton(
            self.btns_flip, text="Next", command=self.to_next, width=115, height=60)
        self.btn_back = ctk.CTkButton(
            self.btns_flip, text="Back", command=self.to_back, width=115, height=60)
        self.btn_back.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_next.pack(side=ctk.RIGHT, padx=5, pady=5)
        self.btns_flip.pack(side=ctk.LEFT, padx=10, pady=5)
        self.btn_suggestion = ctk.CTkButton(
            self.btn_y1, text="Suggestion", command=self.suggest, width=250, height=60)
        self.btn_suggestion.pack(side=ctk.RIGHT, padx=10, pady=5)
        self.btn_y1.pack(padx=10, pady=5)
        # btn_home & btn_save
        self.btn_home = ctk.CTkButton(
            self.btn_y2, text="Home", command=self.master.to_home, width=250, height=60)
        self.btn_home.pack(side=ctk.LEFT, padx=10, pady=5)
        self.btn_save = ctk.CTkButton(
            self.btn_y2, text="Save", command=self.save, width=250, height=60)
        self.btn_save.pack(side=ctk.RIGHT, padx=10, pady=5)
        self.btn_y2.pack(padx=10, pady=5)

    # Flip to next card and save current
    def to_next(self) -> None:
        if self.master.is_empty_string(self.term_txt.get("1.0", tk.END)) or self.master.is_empty_string(self.def_txt.get("1.0", tk.END)):
            messagebox.showinfo("Error", "Please input content for all spaces")
        else:
            # Save the current entries
            self.current.term = self.term_txt.get("1.0", tk.END).rstrip("\n")
            self.current.definition = self.def_txt.get(
                "1.0", tk.END).rstrip("\n")
            self._current_idx += 1
            if self._current_idx >= self.cardset.count:
                # Append a new card if the user is creating a new card
                self.cardset.append(
                    Card(id=-1, term="", definition=""))
            self.current = self.current.next
            # Refresh the page
            self.term_lbl.configure(text="Term {}".format(self._current_idx+1))
            self.def_lbl.configure(
                text="Definition {}".format(self._current_idx+1))
            self.term_txt.delete("1.0", tk.END)
            self.def_txt.delete("1.0", tk.END)
            self.term_txt.insert("1.0", self.current.term)
            self.def_txt.insert("1.0", self.current.definition)

    # Flip to previous card and save current
    def to_back(self) -> None:
        if self._current_idx == 0:
            messagebox.showinfo("Error", "First page - cannot go back.")
        elif self.master.is_empty_string(self.term_txt.get("1.0", tk.END)) or self.master.is_empty_string(self.def_txt.get("1.0", tk.END)):
            self.current = self.current.prev
            self.cardset.delete_at(self._current_idx)
            self._current_idx -= 1
            # Refresh the page
            self.term_lbl.configure(
                text="Term {}".format(self._current_idx + 1))
            self.def_lbl.configure(
                text="Definition {}".format(self._current_idx + 1))
            self.term_txt.delete("1.0", tk.END)
            self.def_txt.delete("1.0", tk.END)
            self.term_txt.insert("1.0", self.current.term)
            self.def_txt.insert("1.0", self.current.definition)
        else:
            # Save the current card
            self.current.term = self.term_txt.get("1.0", tk.END).rstrip("\n")
            self.current.definition = self.def_txt.get("1.0", tk.END).rstrip("\n")
            self.current = self.current.prev
            self._current_idx -= 1
            # Refresh the page
            self.term_lbl.configure(
                text="Term {}".format(self._current_idx + 1))
            self.def_lbl.configure(
                text="Definition {}".format(self._current_idx + 1))
            self.term_txt.delete("1.0", tk.END)
            self.def_txt.delete("1.0", tk.END)
            self.term_txt.insert("1.0", self.current.term)
            self.def_txt.insert("1.0", self.current.definition)

    # Adds the database to card set and returns to HomeFrame
    def save(self) -> None:
        if self.master.is_empty_string(self.term_txt.get("1.0", tk.END)) or self.master.is_empty_string(self.def_txt.get("1.0", tk.END)):
            messagebox.showinfo("Error", "Please input content for all spaces")
        else:
            self.current.term = self.term_txt.get("1.0", tk.END).rstrip("\n")
            self.current.definition = self.def_txt.get("1.0", tk.END).rstrip("\n")
            self.master.add_cardset(self.cardset)
            self.master.print_cardsets()
            self.master.to_home()

    # Pops out DefSuggestWindow, showing all suggested definitions
    def suggest(self) -> None:
        self.current.term = self.term_txt.get("1.0", ctk.END)
        self.current.definition = self.def_txt.get("1.0", ctk.END)
        term = self.current.term
        definition = self.current.definition
        print("Things before the window", term, definition)
        toplv = DefSuggestWindow(self, term, definition)
        toplv.focus()


class DefSuggestWindow(ctk.CTkToplevel):
    def __init__(self, master: tk.Frame, term: str, definition: str, *args):
        super().__init__(master, *args)
        # Initialize window
        self.geometry("700x600")
        self.title("Definitions")

        # Data
        self.term = term
        self.definition = definition
        print("Things", self.term, self.definition)
        self.suggestions = self.get_suggestions()

        # Widgets
        self.def_widgets = []
        self.txt_widgets = []
        self.checkbox_vars = []
        self.confirm_btn = ctk.CTkButton(
            self, text="Copy", command=self.copy_defs)
        self.lbl = ctk.CTkLabel(self, text="No close definition.")
        if self.suggestions == None or len(self.suggestions) == 0:
            self.lbl.pack(padx=5, pady=5)
        else:
            for suggestion in self.suggestions:
                def_frame = ctk.CTkFrame(self)
                self.def_widgets.append(def_frame)
                # Text Widget
                text_widget = TextListWidget(def_frame, text=suggestion)
                self.txt_widgets.append(text_widget)
                # Checkboxes
                var = ctk.BooleanVar()
                var.set(False)
                self.checkbox_vars.append(var)
                checkbox = ctk.CTkSwitch(
                    def_frame, text="", variable=var, onvalue=True, offvalue=False)
                # Packing
                text_widget.pack(side=ctk.LEFT, fill=ctk.X,
                                 expand=True, padx=5, pady=5)
                checkbox.pack(side=ctk.RIGHT, padx=5, pady=5)
                def_frame.pack(padx=5, pady=5)
            # Packing button
            self.confirm_btn.pack(padx=5, pady=10)

    # Gets suggested definitions from Scanner class
    def get_suggestions(self) -> List[str]:
        suggestions = app.suggest_def(
            self.term, self.definition)
        if suggestions == None:
            return None
        else:
            defs = []
            for suggestion in suggestions:
                if suggestion[1] != 0:
                    defs.append(suggestion[0])
                else:
                    break
            return defs

    def copy_defs(self) -> None:
        copied = []
        for i, widget in enumerate(self.txt_widgets):
            if self.checkbox_vars[i].get():
                copied.append(widget.text)
        if len(copied) == 1:
            self.master.current.definition = copied[0]
            self.master.def_txt.delete("1.0", ctk.END)
            self.master.def_txt.insert("1.0", self.master.current.definition)
        elif len(copied) > 0:
            print("Copied array", copied)
            text = ""
            for j in range(len(copied) - 1):
                text = text + copied[j] + ";\n"
            text = text + copied[-1]
            self.master.current.definition = text
            self.master.def_txt.delete("1.0", ctk.END)
            self.master.def_txt.insert("1.0", self.master.current.definition)
        self.destroy()


class DefSuggestEditWindow(DefSuggestWindow):
    def __init__(self, master, term: str, definition: str, entry: ctk.CTkEntry, *args):
        super().__init__(master, term, definition, *args)
        self.entry = entry

    def copy_defs(self):
        copied = []
        for i, widget in enumerate(self.txt_widgets):
            if self.checkbox_vars[i].get():
                copied.append(widget.text)
        if len(copied) == 1:
            self.entry.def_txt.delete(ctk.END)
            self.entry.def_txt.insert(0, copied[0])
        elif len(copied) > 0:
            print("Copied array", copied)
            text = ""
            for j in range(len(copied) - 1):
                text = text + copied[j] + ";\n"
            text = text + copied[-1]
            self.entry.def_txt.delete(ctk.END)
            self.entry.def_txt.insert(0, copied[0])
        self.destroy()


class ImageAddSetFrame(ctk.CTkFrame):
    def __init__(self, master: tk.Frame, title: str, paths: List[str], *args):
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Scanning")
        self.master.geometry("1600x1200")

        self.title = title
        self.paths = paths
        self.dimensions = []  # Tuple: (divider, width, height)

        self.term_color_pixels = ()  # Tuple: (x,y)
        self.def_color_pixels = ()
        self.CANVA_MAX_SIZE = 800
        self.image_idx = 0
        self.card_idx = 0
        # Status of cursor (color picking mode)
        self.pick_mode = -1
        self.term_hsv = ()
        self.def_hsv = ()
        self.x = 0
        self.y = 0

        self.today = date.today()
        tmr = self.today + timedelta(1)
        self.cardset = CardSet(id=-1, title=self.title, last_date=self.today.isoformat(
        ), await_date=tmr.isoformat(), create_date=self.today.isoformat())
        self.cardset.insert(Card(-1, "", ""))
        self.current = self.cardset.head

        for path in paths:
            image = Image.open(path)
            divider, width, height = self.get_adjusted_image_size(image)
            self.dimensions.append((divider, width, height))

        # Left Frame
        self.left_frame = ctk.CTkScrollableFrame(self, width=800, height=1000)
        # Initialize canvas for scanning image
        self.canvas_img = ctk.CTkCanvas(self.left_frame)
        self.image_displayed = None
        self.show_img_canvas(paths[0], self.dimensions[0]
                             [1], self.dimensions[0][2])
        # Define the picking color click action
        self.canvas_img.bind("<Button-1>", self.pick_color)
        # Left Frame - packing
        self.canvas_img.pack(side=ctk.TOP, padx=5, pady=10)

        # Right Frame - Initialization
        self.right_frame = ctk.CTkFrame(self)
        # Color picker - Initialization
        self.color_picker_frame = ctk.CTkFrame(self.right_frame)
        self.btn_pick = ctk.CTkButton(
            self.color_picker_frame, text="Pick Colors", command=self.start_pick_colors)
        self.term_canvas = ctk.CTkCanvas(
            self.color_picker_frame, width=50, height=50)
        self.def_canvas = ctk.CTkCanvas(
            self.color_picker_frame, width=50, height=50)
        # Color picker - Packing
        self.btn_pick.pack(side=ctk.LEFT, padx=5, pady=5)
        self.term_canvas.pack(side=ctk.LEFT, padx=5, pady=5)
        self.def_canvas.pack(side=ctk.LEFT, padx=5, pady=5)
        # Card Viewer - Initialization
        self.card_viewer_frame = ctk.CTkFrame(self.right_frame)
        self.term_lbl = ctk.CTkLabel(
            self.card_viewer_frame, text=f"Term {self.card_idx +1}")
        self.def_lbl = ctk.CTkLabel(
            self.card_viewer_frame, text=f"Definition {self.card_idx + 1}")
        self.term_txt = ctk.CTkTextbox(
            self.card_viewer_frame, width=700, height=100)
        self.def_txt = ctk.CTkTextbox(
            self.card_viewer_frame, width=700, height=100)
        # Color Viewer - Filling color
        self.term_canvas.create_rectangle(
            0, 0, 50, 50, outline="black", width=2, fill="#%02x%02x%02x" % (255, 255, 255))
        self.def_canvas.create_rectangle(
            0, 0, 50, 50, outline="black", width=2, fill="#%02x%02x%02x" % (255, 255, 255))
        # Card Viewer - Packing
        self.term_lbl.pack(side=ctk.TOP, padx=5, pady=5)
        self.term_txt.pack(side=ctk.TOP, padx=5, pady=5)
        self.def_txt.pack(side=ctk.BOTTOM, padx=5, pady=5)
        self.def_lbl.pack(side=ctk.BOTTOM, padx=5, pady=5)
        # Buttons - Initialization
        self.btn_frame = ctk.CTkFrame(self.right_frame)
        self.btn_suggest = ctk.CTkButton(
            self.btn_frame, text="Suggest", command=self.suggest, width=50)
        self.btn_scan = ctk.CTkButton(
            self.btn_frame, text="Scan", command=self.start_scan, width=50)
        self.btn_home = ctk.CTkButton(
            self.btn_frame, text="Home", command=self.master.to_home, width=50)
        self.btn_save = ctk.CTkButton(
            self.btn_frame, text="Save", command=self.save_cardset, width=50)
        # Buttons - Packing
        self.btn_suggest.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_scan.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_home.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_save.pack(side=ctk.LEFT, padx=5, pady=5)

        # Image page flipper - initialization
        self.image_page_flipper = ctk.CTkFrame(self.right_frame)
        self.image_flipper_lbl = ctk.CTkLabel(
            self.image_page_flipper, text="Image")
        self.image_prev_btn = ctk.CTkButton(
            self.image_page_flipper, text="Prev", command=self.to_prev_image)
        self.image_next_btn = ctk.CTkButton(
            self.image_page_flipper, text="Next", command=self.to_next_image)
        self.image_num_vars = ctk.StringVar(value=f"1/{len(self.paths)}")
        self.image_num_lbl = ctk.CTkLabel(
            self.image_page_flipper, textvariable=self.image_num_vars)
        # Image page flipper - packing
        self.image_flipper_lbl.pack(side=ctk.LEFT, padx=5, pady=5)
        self.image_prev_btn.pack(side=ctk.LEFT, padx=5, pady=5)
        self.image_num_lbl.pack(side=ctk.LEFT, padx=5, pady=5)
        self.image_next_btn.pack(side=ctk.RIGHT, padx=5, pady=5)
        # Card page flipper - Initialization
        self.card_page_flipper = ctk.CTkFrame(self.right_frame)
        self.card_flipper_lbl = ctk.CTkLabel(
            self.card_page_flipper, text="Card")
        self.card_prev_btn = ctk.CTkButton(
            self.card_page_flipper, text="Prev", command=self.to_prev_card)
        self.card_num_vars = ctk.StringVar(value=f"1/{self.cardset.count}")
        self.card_num_lbl = ctk.CTkLabel(
            self.card_page_flipper, textvariable=self.card_num_vars)
        self.card_next_btn = ctk.CTkButton(
            self.card_page_flipper, text="Next", command=self.to_next_card)
        # Card page flipper - Packing
        self.card_flipper_lbl.pack(side=ctk.LEFT, padx=5, pady=5)
        self.card_prev_btn.pack(side=ctk.LEFT, padx=5, pady=5)
        self.card_num_lbl.pack(side=ctk.LEFT, padx=5, pady=5)
        self.card_next_btn.pack(side=ctk.RIGHT, padx=5, pady=5)
        # Right Frame - Packing
        self.color_picker_frame.pack(side=ctk.TOP, padx=5, pady=10)
        self.card_viewer_frame.pack(side=ctk.TOP, padx=5, pady=5)
        self.btn_frame.pack(side=ctk.TOP, padx=5, pady=5)
        self.image_page_flipper.pack(side=ctk.TOP, padx=5, pady=5)
        self.card_page_flipper.pack(side=ctk.TOP, padx=5, pady=5)

        # Final packing
        self.left_frame.pack(side=ctk.LEFT, padx=20, pady=5)
        self.right_frame.pack(side=ctk.RIGHT, padx=20, pady=5)

    def get_adjusted_image_size(self, image) -> float:
        width, height = image.size
        if width > self.CANVA_MAX_SIZE or height > self.CANVA_MAX_SIZE:
            divider = max(width, height) / self.CANVA_MAX_SIZE
            width = int(width / divider)
            height = int(height / divider)
        else:
            divider = 1
        return divider, width, height

    # Adjust width and height of image and
    def show_img_canvas(self, path, width, height) -> None:
        image = Image.open(path)
        image = image.resize((width, height))
        self.image_displayed = ImageTk.PhotoImage(image)
        self.canvas_img.configure(width=width, height=height)
        self.canvas_img.create_image(
            0, 0, image=self.image_displayed, anchor="nw")

    # Change appearance of cursor and self.pick_mode
    def start_pick_colors(self) -> None:
        self.configure(cursor="crosshair")
        self.pick_mode = 1  # Adjust to pick term color

    # Return a tuple representing rgb values in int
    def _get_color(self, path, x, y) -> tuple:
        image = Image.open(path)
        rgb = image.getpixel((x, y))
        return rgb

    # Creates a colored square
    def show_color(self, rgb: tuple, canva: ctk.CTkCanvas, size=50) -> None:
        canva.create_rectangle(
            0, 0, size, size, fill="#%02x%02x%02x" % (rgb[0], rgb[1], rgb[2]))

    # Color picker cursor
    def pick_color(self, event) -> None:
        self.y = int(event.y * self.dimensions[self.image_idx][0])
        self.x = int(event.x * self.dimensions[self.image_idx][0])

        if self.pick_mode == 1:  # If just picking term color:
            rgb = self._get_color(
                path=self.paths[self.image_idx], x=self.x, y=self.y)
            self.show_color(rgb, self.term_canvas)
            rgb = np.uint8([[list(rgb)]])
            self.term_hsv = self._cvt_rgb2hsv(rgb)
            self.pick_mode = 2
            self.term_color_pixels = (self.y, self.x)

        elif self.pick_mode == 2:  # If just picking def color:
            rgb = self._get_color(
                path=self.paths[self.image_idx], x=self.x, y=self.y)
            self.show_color(rgb, self.def_canvas)
            rgb = np.uint8([[list(rgb)]])
            self.def_hsv = self._cvt_rgb2hsv(rgb)
            self.pick_mode = -1
            # Adjust cursor appearance
            self.def_color_pixels = (self.y, self.x)
            self.def_hsv = self._cvt_rgb2hsv(rgb)
            self.configure(cursor="arrow")

    def _scan_one_image(self, path) -> None:
        temp = False
        if self.cardset.count == 1:
            temp = True
        terms, defs = self.master.scan_image(
            path, self.term_hsv, self.def_hsv, mask_limit=0.1)
        print(terms, defs)
        for i in range(max(len(terms), len(defs))):
            if i >= len(terms):
                term_card = "<blank>"
            else:
                term_card = terms[i]
            if i >= len(defs):
                def_card = "<blank>"
            else:
                def_card = defs[i]
            self.cardset.append(
                Card(id=-1, term=term_card, definition=def_card))
        if temp:
            # Remove first card
            c = self.cardset.head.next
            head = self.cardset.head
            self.cardset.head = c
            c.prev = None
            del head
        print(self.cardset)

    def scan_images(self, image_idxs: List[int]) -> None:
        for i in image_idxs:
            path = self.paths[i]
            self._scan_one_image(path)
        self.current = self.cardset.head
        self.card_idx = 0
        self.term_txt.delete("1.0", ctk.END)
        self.def_txt.delete("1.0", ctk.END)
        self.term_txt.insert("1.0", text=self.current.term)
        self.def_txt.insert("1.0", text=self.current.definition)
        self.term_lbl.configure(text=f"Term {self.card_idx + 1}")
        self.def_lbl.configure(text=f"Definition {self.card_idx + 1}")
        self.card_num_vars.set(f"{self.card_idx + 1}/{self.cardset.count}")

    def start_scan(self) -> None:
        if len(self.term_color_pixels) == 0 or len(self.def_color_pixels) == 0:
            messagebox.showinfo("Error", "Please select colors first")
        else:
            toplv = SelectImageWindow(self, self.paths)
            toplv.focus()

    def to_next_card(self) -> None:
        term = self.term_txt.get("1.0", ctk.END)
        definition = self.def_txt.get("1.0", ctk.END)
        if self.master.is_empty_string(term) or self.master.is_empty_string(definition):
            messagebox.showinfo("Error", "Empty entry.")
        else:
            if self.current.next == None:
                self.cardset.append(Card(id=-1, term="", definition=""))
            self.current.term = term
            self.current.definition = definition
            self.current = self.current.next
            self.card_idx += 1
            # Refresh widgets
            self.term_lbl.configure(text=f"Term {self.card_idx + 1}")
            self.def_lbl.configure(text=f"Definition {self.card_idx + 1}")
            self.card_num_vars.set(f"{self.card_idx + 1}/{self.cardset.count}")
            self.term_txt.delete("1.0", tk.END)
            self.def_txt.delete("1.0", tk.END)
            self.term_txt.insert("1.0", self.current.term)
            self.def_txt.insert("1.0", self.current.definition)

    def to_prev_card(self) -> None:
        term = self.term_txt.get("1.0", ctk.END)
        definition = self.def_txt.get("1.0", ctk.END)
        if self.master.is_empty_string(term) or self.master.is_empty_string(definition):
            messagebox.showinfo("Error", "Empty entry.")
        elif self.current.prev != None:
            self.current.term = term
            self.current.definition = definition
            self.current = self.current.prev
            self.card_idx -= 1
            # Refresh widgets
            self.term_lbl.configure(text=f"Term {self.card_idx + 1}")
            self.def_lbl.configure(text=f"Definition {self.card_idx + 1}")
            self.card_num_vars.set(
                f"{self.card_idx + 1}/{len(self.cardset.count)}")
            self.term_txt.delete("1.0", tk.END)
            self.def_txt.delete("1.0", tk.END)
            self.term_txt.insert("1.0", self.current.term)
            self.def_txt.insert("1.0", self.current.definition)

    def _cvt_rgb2hsv(self, rgb):
        hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)[0][0]
        return hsv

    def to_prev_image(self):
        print(self.image_idx)
        if self.image_idx - 1 >= 0:
            self.image_idx -= 1
            self.show_img_canvas(
                self.paths[self.image_idx], self.dimensions[self.image_idx][1], self.dimensions[self.image_idx][2])
            self.image_num_vars.set(f"{self.image_idx}/{len(self.paths)}")
        else:
            print("First page")

    def to_next_image(self):
        if self.image_idx + 1 < len(self.paths):
            self.image_idx += 1
            self.show_img_canvas(
                self.paths[self.image_idx], self.dimensions[self.image_idx][1], self.dimensions[self.image_idx][2])
            self.image_num_vars.set(f"{self.image_idx}/{len(self.paths)}")
        else:
            print("Last page")

    def to_home(self):
        self.master.to_home()

    def save_cardset(self):
        current_term = self.term_txt.get("1.0", ctk.END)
        current_def = self.def_txt.get("1.0", ctk.END)
        if self.master.is_empty_string(current_term) or self.master.is_empty_string(current_def):
            messagebox.showinfo(
                "Error", "Please do not leave the boxes blank.")
        else:
            self.master.add_cardset(self.cardset)
            self.master.to_home()

    def suggest(self):
        self.current.term = self.term_txt.get("1.0", ctk.END)
        self.current.definition = self.def_txt.get("1.0", ctk.END)
        term = self.current.term
        definition = self.current.definition
        print("Things before the window", term, definition)
        toplv = DefSuggestWindow(self, term, definition)
        toplv.focus()


class SelectImageWindow(ctk.CTkToplevel):
    def __init__(self, master: tk.Tk, paths: List[str], *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.geometry("1300x800")
        self.title("Select images to scan")

        # Data
        self.paths = paths
        self.images = []
        self.image_check_vars = []

        # Images - Initialization
        self.image_frame = ctk.CTkScrollableFrame(self, width=1200, height=600)
        self.image_canva_frames = []
        # Open images
        for path in paths:
            image = Image.open(path)
            resized_image = image.resize((150, 200))
            resized_image = ImageTk.PhotoImage(resized_image)
            self.images.append(resized_image)
        # Initializing and packing all images and their switches
        for i in range(0, len(self.paths)//4 + 1):
            row_frame = ctk.CTkFrame(self.image_frame)
            for j in range(4):
                if i*4 + j < len(self.paths):
                    # Image Canva widget
                    frame = ctk.CTkFrame(row_frame)
                    canva = ctk.CTkCanvas(frame, width=150, height=200)
                    canva.create_image(
                        0, 0, image=self.images[i*4+j], anchor="nw")
                    canva.create_text(
                        75, 30, text=f"{i*4+j+1}", font=("Helvetica", 24), anchor="center")
                    # Switch widget
                    check_var = ctk.BooleanVar(value=False)
                    switch = ctk.CTkSwitch(
                        frame, onvalue=True, offvalue=False, variable=check_var, text="")
                    canva.pack(side=ctk.TOP, padx=10, pady=5)
                    switch.pack(side=ctk.BOTTOM, padx=10, pady=5)
                    frame.pack(side=ctk.LEFT, padx=10, pady=5)
                    # Save
                    self.image_canva_frames.append(frame)
                    self.image_check_vars.append(check_var)
                else:
                    break
            row_frame.pack(side=ctk.TOP, padx=5, pady=10)

        # Buttons Frame - Initialization
        self.btns_frame = ctk.CTkFrame(self)
        self.back_btn = ctk.CTkButton(
            self.btns_frame, text="Back", command=self.master.master.return_frame)
        self.select_all_btn = ctk.CTkButton(
            self.btns_frame, text="Select All", command=self.select_all)
        self.clear_btn = ctk.CTkButton(
            self.btns_frame, text="Clear Choices", command=self.clear_all)
        self.done_btn = ctk.CTkButton(
            self.btns_frame, text="Done", command=self.done)
        # Buttons Frame - Packing
        self.back_btn.pack(side=ctk.LEFT, padx=5, pady=5)
        self.select_all_btn.pack(side=ctk.LEFT, padx=5, pady=5)
        self.clear_btn.pack(side=ctk.LEFT, padx=5, pady=5)
        self.done_btn.pack(side=ctk.LEFT, padx=5, pady=5)

        # Final Packing
        self.image_frame.pack(side=ctk.TOP, padx=5, pady=10)
        self.btns_frame.pack(side=ctk.BOTTOM, padx=5, pady=10)

    def select_all(self):
        for var in self.image_check_vars:
            var.set(True)

    def clear_all(self):
        for var in self.image_check_vars:
            var.set(False)

    def done(self):
        idxs = []
        for i, var in enumerate(self.image_check_vars):
            if var:
                idxs.append(i)
        self.master.scan_images(idxs)
        self.destroy()


class DocumentChoiceFrame(ctk.CTkFrame):
    def __init__(self, master, title, *args) -> None:
        super().__init__(master, *args)
        # Initialize window size
        self.master.title("Import by...")
        self.master.geometry("210x135")

        # Data
        self.title = title

        # Widgets (buttons)
        self.btn_text = ctk.CTkButton(
            self, text="Text", command=self.import_text)
        self.btn_excel = ctk.CTkButton(
            self, text="Excel", command=self.import_excel)
        self.btn_word = ctk.CTkButton(
            self, text="Word", command=self.import_word)

        self.btn_text.pack(pady=5)
        self.btn_excel.pack(pady=5)
        self.btn_word.pack(pady=5)

        self.cardset = CardSet(id=-1, title=self.title)

    def import_text(self):
        self.master.open_frame(TextAddSetFrame, self.title)

    # Imports a card set by reading excel via pandas
    def import_excel(self) -> None:
        self.filename = filedialog.askopenfilename(
            filetypes=[("Excel", "*.xlsx")])
        if self.filename != "":
            dataframe = pd.read_excel(self.filename, engine="openpyxl")
            terms = dataframe["terms"]
            defs = dataframe["defs"]
            for i in range(len(terms)):
                self.cardset.append(Card(
                    id=-1, term=terms[i], definition=defs[i]))
                self.master.add_cardset(self.cardset)
                self.master.to_home()
            self.master.to_home()

    # Imports a card set by reading word document with input file format
    def import_word(self) -> None:
        # Ask for file directory
        self.filename = filedialog.askopenfilename(
            filetypes=[("Word documents", "*.docx")])
        # If the user selects a file
        if self.filename != "":
            doc = docx.Document(self.filename)
            for paragraph in doc.paragraphs:
                texts = paragraph.text.split(";; ", maxsplit=1)
                self.cardset.append(Card(id=-1, term=texts[0], definition=texts[1]))
        self.master.add_cardset(self.cardset)
        self.master.to_home()


class TextAddSetFrame(ctk.CTkFrame):
    def __init__(self, master, title, *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Add by...")
        self.master.geometry("700x600")

        # Initialize data
        self.today = date.today()
        tmr = self.today + timedelta(1)
        self.cardset = CardSet(id=-1, title=title, last_date=self.today.isoformat(),
                               await_date=tmr.isoformat(), create_date=self.today.isoformat())
        self.title = title

        # Widgets
        self.lbl_text = ctk.CTkLabel(self, text="Text")
        self.lbl_pattern = ctk.CTkLabel(
            self, text="Pattern (represent term by {term} and definition by {def}), and no need to input backspace")
        font = ctk.CTkFont(family="Helvetica", size=20)
        self.txt_text = ctk.CTkTextbox(self, width=600, font=font)
        self.txt_pattern = ctk.CTkTextbox(self, width=600, font=font)

        self.lbl_text.pack(padx=5, pady=5)
        self.txt_text.pack(padx=5, pady=5)
        self.lbl_pattern.pack(padx=5, pady=5)
        self.txt_pattern.pack(padx=5, pady=5)

        self.btns = ctk.CTkFrame(self)
        self.btn_home = ctk.CTkButton(
            self.btns, text="Home", command=self.to_home)
        self.btn_home.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_import = ctk.CTkButton(
            self.btns, text="Import", command=self.import_cardset)
        self.btn_import.pack(side=ctk.RIGHT, padx=5, pady=5)
        self.btns.pack(padx=5, pady=5)

    # Returns to HomeFrame without saving
    def to_home(self) -> None:
        self.master.to_home()

    # Saves the new card set into database
    def import_cardset(self) -> None:
        string = self.txt_text.get("1.0", tk.END)
        pattern = self.txt_pattern.get("1.0", tk.END)
        # If the user did not input a pattern
        if self.master.is_empty_string(pattern):
            pattern = "{term};; {def}"
        self._read_txt_pattern(string, pattern)
        self.master.add_cardset(self.cardset)
        self.to_home()

    # Read the pattern and text and appends them into self.cardset
    def _read_txt_pattern(self, string: str, pattern: str) -> None:
        delimiter1 = "{term}"
        delimiter2 = "{def}"
        # Get the pattern indicators
        if not (delimiter1 in pattern and delimiter2 in pattern):
            messagebox.showinfo(
                "Error!", "Please ensure you inputted {term} and {def}")
            return
        else:
            substrings = pattern.split(delimiter1)
            if len(substrings) > 2:
                messagebox.showinfo(
                    "Error!", "Please ensure you only input {term} once.")
                return
            seg1, seg2 = substrings[0], substrings[1]
            substrings = seg2.split(delimiter2)
            if len(substrings) > 2:
                messagebox.showinfo(
                    "Error!", "Please ensure that you only input {def} once.")
                return
            seg2, seg3 = substrings[0], substrings[1]
        if self.master.is_empty_string(seg3):
            print("Empty seg3")
            seg3 = "\n"
        else:
            seg3 = seg3 + "\n"
        print(f"Segment 1 '{seg1}', Segment 2 '{seg2}', Segment 3 '{seg3}'")
        # Get the cardsets via the segments
        substrings = string.split(seg2)
        terms = []
        defs = []
        # Split the first term
        if seg1 != "":
            terms.append(substrings[0].split(seg1, maxsplit=1)[-1])
        else:
            terms.append(substrings[0])
        # Split all middle terms
        for i in range(1, len(substrings)-1):
            sub = substrings[i].split(seg3, maxsplit=1)
            print("Check boolean", seg3 == "\n")
            print("Sub", sub)
            defs.append(sub[0])
            terms.append(sub[1])
            print("Info", defs[-1], terms[-2])
        # Split the last def
        defs.append(substrings[0].split(seg3, maxsplit=3)[0])

        # Get them to self.cardset
        for i in range(len(terms)):
            new_card = Card(
                id=-1, term=terms[i], definition=defs[i])
            self.cardset.append(new_card)


class WordAddSetFrame(ctk.CTkFrame):
    def __init__(self, master, title, *args) -> None:
        super().__init__(master, *args)
        # Initialize window
        self.master.title("Import by Word")
        self.master.geometry("700x600")

        # Initialize data
        self.title = title
        self.today = date.today()
        tmr = self.today + timedelta(1)
        self.cardset = CardSet(id=-1, title=title)

        # Widgets
        font = tk.font.Font(family="Helvetica", size=20)
        self.lbl_info = tk.Label(
            self, text="The word document should repesent each term and word in a line. Input the pattern for each line below.")
        self.lbl_pattern = tk.Label(
            self, text="Pattern (represent term by {term} and definition by {def}), and no need to input backspace")
        self.txt_pattern = self.txt_pattern = tk.Text(
            self, height=3, font=font)
        self.btns = tk.Frame(self)
        self.btn_home = tk.Button(self.btns, text="Home", command=self.to_home)
        self.btn_import = tk.Button(
            self.btns, text="Import", command=self.import_word)
        self.btn_home.pack(side=tk.LEFT, padx=5, pady=5)
        self.btn_import.pack(side=tk.RIGHT, padx=5, pady=5)

        self.lbl_info.pack(padx=5, pady=5)
        self.lbl_pattern.pack(padx=5, pady=5)
        self.txt_pattern.pack(padx=5, pady=5)
        self.btns.pack(padx=5, pady=5)

    # Returns to home page
    def to_home(self) -> None:
        self.master.to_home()

    def import_word(self) -> None:
        self.filename = filedialog.askopenfilename(
            filetypes=[("Word documents", "*.docx")])
        # If the user selected a file
        if self.filename != "":
            doc = docx.Document(self.filename)
            texts = []
            for paragraph in doc.paragraphs:
                texts.append(paragraph.text)
            pattern = self.txt_pattern.get("1.0", tk.END)
            if self.master.is_empty_string(pattern):
                pattern = "{term};; {def}"
            seg1, seg2, seg3 = self._read_pattern(pattern)
            self._read_text(texts, seg1, seg2, seg3)
            self.master.add_cardset(self.cardset)
            self.master.to_home()

    ''' TODO: code assumed {term} before {def} '''

    def _read_text(self, paras: List[str], seg1: str, seg2: str, seg3: str) -> None:
        for para in paras:
            substrings = para.split(seg2, maxsplit=1)
            self.cardset.append(
                Card(id=-1, term=substrings[0], definition=substrings[1]))

    def _read_pattern(self, pattern: str) -> str:
        delimiter1 = "{term}"
        delimiter2 = "{def}"
        if not (delimiter1 in pattern and delimiter2 in pattern):
            raise ValueError("Pattern does not have delimiters.")
        substrings = pattern.split(delimiter1, maxsplit=1)
        seg1, pattern = substrings[0], substrings[1]
        substrings = pattern.split(delimiter2, maxsplit=1)
        seg2, seg3 = substrings[0], substrings[1]
        return seg1, seg2, seg3

    def _read_txt_pattern(self, string: str, pattern: str) -> None:
        delimiter1 = "{term}"
        delimiter2 = "{def}"
        # Get the pattern indicators
        if not (delimiter1 in pattern and delimiter2 in pattern):
            messagebox.showinfo(
                "Error!", "Please ensure you inputted {term} and {def}")
            return
        else:
            substrings = pattern.split(delimiter1)
            if len(substrings) > 2:
                messagebox.showinfo(
                    "Error!", "Please ensure you only input {term} once.")
                return
            seg1, seg2 = substrings[0], substrings[1]
            substrings = seg2.split(delimiter2)
            if len(substrings) > 2:
                messagebox.showinfo(
                    "Error!", "Please ensure that you only input {def} once.")
                return
            seg2, seg3 = substrings[0], substrings[1]
        if self.master.is_empty_string(seg3):
            print("Empty seg3")
            seg3 = "\n"
        else:
            seg3 = seg3 + "\n"
        print(f"Segment 1 '{seg1}', Segment 2 '{seg2}', Segment 3 '{seg3}'")
        # Get the cardsets via the segments
        substrings = string.split(seg2)
        terms = []
        defs = []
        # Split the first term
        if seg1 != "":
            terms.append(substrings[0].split(seg1, maxsplit=1)[-1])
        else:
            terms.append(substrings[0])
        # Split all middle terms
        for i in range(1, len(substrings)-1):
            sub = substrings[i].split(seg3, maxsplit=1)
            print("Check boolean", seg3 == "\n")
            print("Sub", sub)
            defs.append(sub[0])
            terms.append(sub[1])
            print("Info", defs[-1], terms[-2])
        # Split the last def
        defs.append(substrings[0].split(seg3, maxsplit=3)[0])

        # Get them to self.cardset
        for i in range(len(terms)):
            new_card = Card(
                id=-1, term=terms[i], definition=defs[i])
            self.cardset.append(new_card)


class CardFrame(ctk.CTkFrame):
    def __init__(self, master, cardset_id: int, *args):
        super().__init__(master, *args)
        # Get information
        self.cardset_id = cardset_id
        self.cardset = self.master.get_cardset(cardset_id)
        self.cardset = self.master.get_set_cards(
            self.cardset, self.cardset_id, mode="default")
        self.no_new = True
        self.width = 1300
        self.height = 600
        self.today = date.today().isoformat()

        # Initialize window
        self.master.title(self.cardset.title)
        self.master.geometry("1300x600")

        # Widgets (buttons)
        self.btns = ctk.CTkFrame(self, border_color="dark_color", width=1300)
        self.btn_back = ctk.CTkButton(
            self.btns, text="Back", command=self.master.to_home, width=250, height=60)
        self.btn_edit = ctk.CTkButton(
            self.btns, text="Edit", command=self.start_edit, width=250, height=60)
        self.btn_del = ctk.CTkButton(
            self.btns, text="Delete", command=self.start_delete, width=250, height=60)
        self.btn_memor = ctk.CTkButton(
            self.btns, text="Memorize", command=self.start_memorize, width=250, height=60)
        # Buttons that are visible only when editing
        self.btns_edit_frame = ctk.CTkFrame(self)
        self.btn_add = ctk.CTkButton(
            self.btns_edit_frame, text="Add card", command=self.add_entry)
        self.btn_suggest = ctk.CTkButton(
            self.btns_edit_frame, text="Suggest", command=self.suggest)
        # Pack these buttons
        self.btn_add.pack(side=ctk.LEFT, padx=5, pady=5)
        self.btn_suggest.pack(side=ctk.RIGHT, padx=5, pady=5)

        # Pack the buttons
        self.btn_back.pack(side=ctk.LEFT, pady=5, padx=10)
        self.btn_edit.pack(side=ctk.LEFT, fill=ctk.X,
                           expand=True, pady=5, padx=10)
        self.btn_memor.pack(side=ctk.RIGHT, fill=ctk.X,
                            expand=True, pady=5, padx=10)
        self.btn_del.pack(side=ctk.RIGHT, fill=ctk.X,
                          expand=True, pady=5, padx=10)

        self.btns.pack(side=ctk.TOP, pady=5, padx=10)

        self.checkList = []  # Stores all the variable pointers of the checkbuttons
        # Show the words
        self.words_frame = ctk.CTkScrollableFrame(
            self, border_color="dark_color", width=1300, height=450)
        # Widgets for default cards
        self.widget_words = []
        # Widgets for entry cards
        self.widget_entries = []
        current = self.cardset.head
        while current.next != None:
            del_var = ctk.BooleanVar()
            word = WordDelWidget(self.words_frame, current, del_var)
            word.pack(padx=5, pady=5, fill=ctk.X, expand=True)
            self.widget_words.append(word)
            current = current.next
        del_var = ctk.BooleanVar()
        word = WordDelWidget(self.words_frame, current, del_var)
        word.pack(side=ctk.TOP, padx=5, pady=5, fill=ctk.X, expand=True)
        self.widget_words.append(word)
        self.words_frame.pack(side=ctk.TOP, padx=10, pady=5)

    def start_memorize(self) -> None:
        print("Cardset to memorize:")
        print(self.cardset)
        current = self.cardset.head
        print("Initializing...")
        cardset = CardSet(
            self.cardset.id, self.cardset.title, self.cardset.last_date, self.cardset.await_date, self.cardset.CREATE_DATE)
        while current.next != None and current.next.memor_date <= self.today:
            print(current)
            cardset.random_insert(current, 0, cardset.count)
            print("Count", self.cardset.count)
            current = current.next
        if current.memor_date > self.today:
            messagebox.showinfo("Reminder", "No cards to memorize")
        else:
            print(current)
            print("==========")
            cardset.random_insert(current, 0, cardset.count)
            self.master.open_frame(MemorizeFrame, cardset)

    def back2Home(self):
        self.master.config(menu=tk.Menu(self.master))
        self.master.to_home()

    # Shows all delete checkboxes if the card set has more than one word
    def start_delete(self) -> None:
        if len(self.widget_words) == 1:
            messagebox.showinfo(
                "Reminder", "You can't delete all words in a cardset!")
        else:
            for widget in self.widget_words:
                widget.show_checkboxes()
            self.btn_del.configure(text="Done", command=self.save_delete)

    # Saves updates of deletion into database and refreshes GUI
    def save_delete(self) -> None:
        del_id = []
        for widget in self.widget_words:
            if widget.del_var.get():
                del_id.append(widget.card_id)
        if len(del_id) == len(self.widget_words):
            messagebox.showinfo(
                "Reminder", "You can't delete all words in a cardset!")
        else:
            self.master.delete_cards_id(del_id)
            self.master.set_last_date(self.cardset_id, self.today)
            self.master.refresh_frame(CardFrame, self.cardset_id)

    # Shows entry widgets
    def start_edit(self) -> None:
        for i, widget in enumerate(self.widget_words):
            widget.pack_forget()
            self.widget_entries.append(
                WordEntryWidget(self.words_frame, widget.card))
            self.widget_entries[i].pack()
        self.btn_edit.configure(text="Done", command=self.save_edit)
        self.btns_edit_frame.pack(padx=5, pady=10)

    # Save edit results into database and shows default cards widgets
    def save_edit(self) -> None:
        current = self.cardset.head
        lasted = False
        for i, widget in enumerate(self.widget_entries):
            if self.master.is_empty_string(widget.term_txt.get()) or self.master.is_empty_string(widget.def_txt.get()):
                messagebox.showinfo("Error", "You should not leave blanks.")
            else:
                term = widget.term_txt.get()
                definition = widget.def_txt.get()
                print("Data", term, definition)
                if self.no_new:
                    if current.next != None:
                        current.term = term
                        current.definition = definition
                        self.master.set_card(current.id, current)
                        current = current.next
                    else:
                        current.term = term
                        current.definition = definition
                        self.master.set_card(current.id, current)
                # Add new card
                else:
                    print("Adding...")
                    if current.next != None:
                        print("First few cards...")
                        current.term = term
                        current.definition = definition
                        self.master.set_card(current.id, current)
                        current = current.next
                    elif lasted:
                        print("Adding new card...")
                        c = Card(-1, term, definition)
                        print(c)
                        self.master.print_cards()
                        self.master.add_card(c, self.cardset_id)
                        self.master.print_cards()
                    else:
                        print("Updating last preexisting card")
                        # The last card originally in database
                        current.term = term
                        current.definition = definition
                        self.master.set_card(current.id, current)
                        lasted = True
        self.master.set_last_date(self.cardset_id, self.today)
        self.master.refresh_frame(CardFrame, self.cardset_id)

    def add_entry(self) -> None:
        self.widget_entries.append(WordEntryWidget(
            self.words_frame, Card(-1, "", "")))
        self.widget_entries[-1].pack()
        self.no_new = False

    # Show Toplevel for suggesting definitions
    def suggest(self) -> None:
        focus = self.master.focus_get()
        if focus:
            master = focus.master.master
            term = master.term_txt.get()
            definition = master.def_txt.get()
            print("Things before the window", term, definition)
            toplv = DefSuggestEditWindow(self, term, definition, master)
            toplv.focus()
        else:
            messagebox.showinfo(
                "Error", "Make sure that you are focusing on a text box.")


class MemorizeFrame(ctk.CTkFrame):
    def __init__(self, master, cardset: CardSet, *args) -> None:
        super().__init__(master, *args)
        # Initialize size of window
        self.master.title(cardset.title)
        self.master.geometry("1300x600")

        # Initialize data
        self.cardset = cardset
        self.current = self.cardset.head
        self.once_count = 0
        self.finished_count = 0
        self.TOTAL_COUNT = self.cardset.count
        self.today = date.today()

        # Widgets
        # Widgets on left
        font = ctk.CTkFont(family="Helvetica", size=20)
        self.left = ctk.CTkFrame(self)
        self.txt_term = ctk.CTkTextbox(
            master=self.left, font=font, height=200, width=1000)
        self.txt_def = ctk.CTkTextbox(
            master=self.left, font=font, height=300, width=1000)
        self.txt_term.insert("1.0", self.current.term)
        self.txt_def.insert("1.0", f"Definition")
        self.txt_term.configure(state="disabled")
        self.txt_def.configure(state="disabled", fg_color="gray")
        self.progress_val = self.finished_count / self.TOTAL_COUNT
        self.progress = ttk.Progressbar(
            self.left, orient="horizontal", length=720, mode="determinate", variable=self.progress_val)
        self.txt_term.pack(padx=5, pady=10)
        self.txt_def.pack(padx=5, pady=10)
        self.progress.pack(padx=5, pady=10)
        # Widgets on right
        self.right = ctk.CTkFrame(self)
        self.btn_back = ctk.CTkButton(
            self.right, text="Back", command=self.back, height=60, width=250)
        self.btn_show = ctk.CTkButton(
            self.right, text="Show", command=self.show, height=60, width=250)
        self.btn_forgot = ctk.CTkButton(
            self.right, text="Forgot", command=self.forgot, height=60, width=250)
        self.btn_memorize = ctk.CTkButton(
            self.right, text="Memorize", command=self.memorize, height=60, width=250)
        self.btn_back.pack(padx=5, pady=20)
        self.btn_show.pack(padx=5, pady=20)
        self.btn_forgot.pack(padx=5, pady=20)
        self.btn_memorize.pack(padx=5, pady=20)

        # Pack the rest of widgets
        self.left.pack(side=tk.LEFT, padx=10, pady=5)
        self.right.pack(side=tk.RIGHT, padx=10, pady=5)

    # Returns to CardFrame
    def back(self) -> None:
        self.master.destroy_and_return(self.cardset.id)

    # Change the display so that a new card is shown
    def change_displayed_card(self, front: str, back: str) -> None:
        self.txt_term.configure(state="normal")
        self.txt_def.configure(state="normal", fg_color="gray")
        self.txt_term.delete("1.0", tk.END)
        self.txt_def.delete("1.0", tk.END)
        self.txt_term.insert("1.0", front)
        self.txt_def.insert("1.0", back)
        self.txt_term.configure(state="disabled")
        self.txt_def.configure(state="disabled")
        self.progress_val = self.finished_count / self.TOTAL_COUNT

    def show(self) -> None:
        self.txt_def.configure(state="normal", fg_color="white")
        self.txt_def.delete("1.0", tk.END)
        self.txt_def.insert("1.0", self.current.definition)
        self.txt_def.configure(state="disabled")

    def forgot(self) -> None:
        self.current.forgot = True
        if self.current.state >= 2:
            raise ValueError(
                "State is equal or greater than 2 when clicking Forgot")
        elif self.current.state == 1:
            self.once_count -= 1
            self.current.state = 0
        # If current is the last card
        if self.current.next == None:
            print("Deleting the last card...")
            last_card = self.current
            self.cardset.delete_at(-1)
            self.cardset.random_insert(
                last_card, 0, self.cardset.count - self.once_count)
            self.current = self.cardset.head
            self.change_displayed_card(
                self.current.term, self.current.definition)
            print(self.cardset)
        else:
            print("Deleting card before last...")
            next_card = self.current.next
            self.cardset.delete(self.current)
            self.cardset.random_insert(
                self.current, 0, self.cardset.count - self.once_count)
            self.current = next_card
            self.change_displayed_card(
                self.current.term, self.current.definition)
            print(self.cardset)

    def memorize(self) -> None:
        self.current.state += 1
        print(f"Current state: {self.current.state}")
        if self.current.state > 2 or self.current.state <= 0:
            raise ValueError("State invalid!")
        elif self.current.state == 2:
            # Finish memorization
            if self.cardset.count == 1:
                self._update_box()
                self.master.set_last_date(
                    self.cardset.id, self.today.isoformat())
                self.master.set_await_date(
                    self.cardset.id, self.cardset.await_date)
                # Updates cards, cardset and
                messagebox.showinfo("Congratulations!",
                                    "You finished memorizing!")
                self.back()
            else:
                self._update_box()
                # Update the current card's status
                self.once_count -= 1
                # If current is last card
                if self.current.next == None:
                    print("Deleting last card...")
                    last_card = self.current
                    self.cardset.delete(last_card)
                    self.current = self.cardset.head
                    self.change_displayed_card(
                        self.current.term, self.current.definition)
                    print(self.cardset)
                else:
                    print("Deleting card before last...")
                    next_card = self.current.next
                    self.cardset.delete(self.current)
                    self.current = next_card
                    self.change_displayed_card(
                        self.current.term, self.current.definition)
                    print(self.cardset)
        # Moves to right of linked list
        else:
            # If current is last card
            if self.current.next == None:
                print("Deleting last card...")
                last_card = self.current
                self.cardset.delete_at(-1)
                self.cardset.random_insert(
                    last_card, self.cardset.count - self.once_count, self.cardset.count)
                self.once_count += 1
                self.current = self.cardset.head
                self.change_displayed_card(
                    self.current.term, self.current.definition)
                print(self.cardset)
                print(f"Once count {self.once_count}")
            else:
                print("Deleting card before last card...")
                next_card = self.current.next
                self.cardset.delete(self.current)
                self.cardset.random_insert(
                    self.current, self.cardset.count - self.once_count, self.cardset.count)
                self.once_count += 1
                self.current = next_card
                self.change_displayed_card(
                    self.current.term, self.current.definition)
                print(self.cardset)
                print(f"Once count: {self.once_count}")

    def _update_box(self) -> None:
        if self.current.forgot:
            self.current.box = 0
            self.master.set_box(
                self.current.id, self.current.box, self.cardset)
        else:
            self.current.box += 1
            self.master.set_box(
                self.current.id, self.current.box, self.cardset)
        self._update_date()

    def _update_date(self) -> None:
        # Update cards date
        if self.current.box < len(self.master.leitner_pattern):
            next_day = self.today + \
                timedelta(self.master.leitner_pattern[self.current.box])
            next_day = next_day.isoformat()
            self.current.memor_date = next_day
            # Update card set await date
            if self.cardset.await_date < next_day:
                self.cardset.await_date = next_day
        # TODO: when the box is over... '''
        elif self.current.box == 8:
            print("Yay")
        else:
            self.current.finished_memor = True

# Main Application
class FlashcardApp(ctk.CTk, Database, Scanner):
    def __init__(self) -> None:
        super().__init__()
        Database.__init__(self)
        Scanner.__init__(self)
        self.cardsets = self.get_all_cardsets()
        home_frame = HomeFrame(master=self, cardsets=self.cardsets)
        self.frame_stack = [home_frame]
        self.frame_stack[0].pack()
        self.current_idx = 0
        ''' TODO: change leitner pattern to more reasonable '''
        self.leitner_pattern = [1, 3, 5]
        # Sets the end connection function
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    # Returns to HomeFrame and destroys all frames
    def to_home(self) -> None:
        self.frame_stack[self.current_idx].destroy()
        del self.frame_stack[:]
        self.cardsets = self.get_all_cardsets()
        home_frame = HomeFrame(self, cardsets=self.cardsets)
        self.frame_stack = [home_frame]
        self.current_idx = 0
        gc.collect()
        self.frame_stack[self.current_idx].pack()

    # Opens a new frame and pushes the frame into stack
    def open_frame(self, frame_class, *args, **kwargs) -> None:
        self.frame_stack[self.current_idx].pack_forget()
        self.frame_stack.append(frame_class(self, *args, **kwargs))
        self.current_idx += 1
        self.frame_stack[self.current_idx].pack(fill=ctk.BOTH, expand=True)

    # Returns to previous frame (as it was) and removes the last inside frame
    def return_frame(self) -> None:
        if self.current_idx > 0:
            self.frame_stack[self.current_idx].pack_forget()
            del self.frame_stack[self.current_idx]
            self.current_idx -= 1
            self.frame_stack[self.current_idx].pack()
        else:
            raise RuntimeError("This is the first page in stack.")

    # Returns to previous frame (by refreshing it) and removes the last inside frame
    def destroy_and_return(self, *args, **kwargs) -> None:
        if self.current_idx > 0:
            self.frame_stack[self.current_idx].pack_forget()
            del self.frame_stack[self.current_idx]
            self.current_idx -= 1
            frame_class = type(self.frame_stack[self.current_idx])
            self.frame_stack[self.current_idx] = frame_class(
                self, *args, **kwargs)
            self.frame_stack[self.current_idx].pack()
        else:
            raise RuntimeError("This is the first page in stack.")

    # Repacks the current frame
    def refresh_frame(self, frame_class: tk.Frame, *args) -> None:
        self.frame_stack[self.current_idx].destroy()
        self.frame_stack[self.current_idx] = frame_class(self, *args)
        self.frame_stack[self.current_idx].pack()

    # Checks whether the string is empty
    @staticmethod
    def is_empty_string(string: str) -> bool:
        string = string.replace(" ", "")
        string = string.replace("\b", "")
        string = string.replace("\n", "")
        if string == "":
            return True
        else:
            return False

    # Define when the user clicks the close button
    def on_closing(self) -> None:
        self.end_connection()
        self.destroy()


if __name__ == "__main__":
    app = FlashcardApp()
    app.mainloop()
